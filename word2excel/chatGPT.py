import docx
import os
import openpyxl
from docx.shared import Inches

# 打开Word文档
# doc = docx.Document('your_file.docx')
doc = docx.Document('701范本.docx')

# 创建Excel工作簿
wb = openpyxl.Workbook()
ws = wb.active

# 添加Excel表头
ws.cell(row=1, column=1, value='标题')
ws.cell(row=1, column=2, value='内容')

# 定义表格图片输出目录
image_folder = 'table_images'

# 遍历文档中的所有段落
for i, para in enumerate(doc.paragraphs):
    # 提取段落中的文本内容
    text = para.text
    
    # 根据标题格式（比如1.1、2.2.1等）判断是否为标题，如果是则将标题和内容写入Excel表格
    if para.style.name.startswith('Heading'):
        ws.cell(row=i+2, column=1, value=text)
    else:
        ws.cell(row=i+2, column=2, value=text)

# 遍历文档中的所有表格
for i, table in enumerate(doc.tables):
    # 将表格保存为图片
    image_path = os.path.join(image_folder, f'table_{i}.png')
    table._element.clear_info()
    table._element.addprevious(table._element.makeelement(docx.oxml.ns.qn('w', 'tblPr')))
    doc.save('temp.docx')
    doc2 = docx.Document('temp.docx')
    doc2.save('temp2.docx')
    doc3 = docx.Document('temp2.docx')
    doc3.save('temp3.docx')
    table_png = doc3.tables[i]._tbl._element.drawings_part.blob
    with open(image_path, 'wb') as f:
        f.write(table_png)

# 保存Excel表格
wb.save('output.xlsx')
